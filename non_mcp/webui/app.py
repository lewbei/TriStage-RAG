#!/usr/bin/env python3
"""
Simple HTML GUI for the TriStage Retrieval Pipeline (non-MCP)

- Loads the non_mcp pipeline using AppConfig and the same stage classes from tristage_rag/
- Provides:
  - HTML search page at /
  - JSON search API at /api/search
  - Endpoints to add documents from text form or file upload

This does not modify tristage_rag; it reuses the non_mcp ThreeStageRetrievalSystem directly.
"""
import os
import sys
from pathlib import Path
import logging
from typing import Optional
import hashlib
import json
import time

from flask import Flask, render_template, request, jsonify, redirect, url_for, flash, session

# Paths
APP_DIR = Path(__file__).parent.resolve()
NON_MCP_DIR = APP_DIR.parent.resolve()
PROJECT_ROOT = NON_MCP_DIR.parent.resolve()

# Allow running directly without `pip install -e .`; no-op once installed.
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

# Import the non_mcp system
from non_mcp.main import AppConfig, ThreeStageRetrievalSystem
# No LLM generation; chat composes answers from retrieved snippets only


def create_app() -> Flask:
    app = Flask(
        __name__,
        template_folder=str(APP_DIR / "templates"),
        static_folder=str(APP_DIR / "static"),
    )
    app.secret_key = os.environ.get("NON_MCP_WEBUI_SECRET", "dev-secret")

    # Logging to file + console
    logging.basicConfig(
        level=getattr(logging, os.environ.get("NON_MCP_WEBUI_LOG_LEVEL", "INFO"), logging.INFO),
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(os.environ.get("NON_MCP_WEBUI_LOG_FILE", "non_mcp_webui.log")),
            logging.StreamHandler()
        ]
    )

    # Build config (align with repo-level defaults)
    config = AppConfig(
        models_dir=str(PROJECT_ROOT / "models"),
        data_dir=str(PROJECT_ROOT / "data"),
        index_dir=str(PROJECT_ROOT / "faiss_index"),
        device=os.environ.get("NON_MCP_DEVICE", "auto"),
        log_level=os.environ.get("NON_MCP_LOG_LEVEL", "INFO"),
    )

    # Initialize system
    system = ThreeStageRetrievalSystem(config)
    # Note: No generator initialized; results are returned as snippets only

    # ---- Helpers ----
    MANIFEST_PATH = PROJECT_ROOT / "data" / "embedded_manifest.json"

    def _ensure_data_dir():
        (PROJECT_ROOT / "data").mkdir(exist_ok=True)

    def load_manifest():
        _ensure_data_dir()
        if MANIFEST_PATH.exists():
            try:
                return json.loads(MANIFEST_PATH.read_text(encoding='utf-8'))
            except Exception:
                app.logger.warning("Failed to read manifest; recreating")
        return {"version": 1, "files": []}

    def save_manifest(m):
        _ensure_data_dir()
        MANIFEST_PATH.write_text(json.dumps(m, indent=2), encoding='utf-8')

    def text_hash(text: str) -> str:
        return hashlib.sha256(text.encode('utf-8')).hexdigest()

    def repo_documents_status():
        repo_docs = PROJECT_ROOT / "documents"
        statuses = []
        if not repo_docs.exists():
            return statuses
        manifest = load_manifest()
        embedded_hashes = {f.get('hash') for f in manifest.get('files', [])}
        for ext in ("*.txt", "*.md", "*.markdown", "*.json", "*.pdf", "*.docx"):
            for p in repo_docs.rglob(ext):
                txt = extract_text_from_path(p)
                h = text_hash(txt) if txt else None
                statuses.append({
                    "path": str(p.relative_to(PROJECT_ROOT)),
                    "hash": h,
                    "embedded": bool(h and h in embedded_hashes),
                    "size_bytes": len(txt.encode('utf-8', errors='ignore')) if txt else 0,
                })
        return statuses
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200):
        text = text.strip()
        if not text:
            return []
        chunks = []
        start = 0
        n = len(text)
        while start < n:
            end = min(start + chunk_size, n)
            # try to end at sentence boundary
            window = text[start:end]
            if end < n:
                last_period = window.rfind('.')
                last_newline = window.rfind('\n')
                cut = max(last_period, last_newline)
                if cut > 0 and (end - (start + cut)) < 200:
                    end = start + cut + 1
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            if end >= n:
                break
            start = max(end - overlap, 0)
        return chunks

    def extract_text_from_upload(fs) -> str:
        """Extract text from uploaded FileStorage supporting txt, md, pdf, docx."""
        filename = (fs.filename or '').lower()
        data = fs.read()
        # reset pointer if needed later
        try:
            fs.stream.seek(0)
        except Exception:
            pass
        if filename.endswith(('.txt', '.md', '.markdown')):
            try:
                return data.decode('utf-8', errors='ignore')
            except Exception:
                return ''
        if filename.endswith('.pdf'):
            try:
                from pypdf import PdfReader
                import io
                reader = PdfReader(io.BytesIO(data))
                pages = [p.extract_text() or '' for p in reader.pages]
                return '\n\n'.join(pages)
            except Exception as e:
                app.logger.exception('PDF parse failed')
                return ''
        if filename.endswith('.docx'):
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(data))
                paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
                return '\n'.join(paras)
            except Exception:
                app.logger.exception('DOCX parse failed')
                return ''
        # Fallback: treat as text
        try:
            return data.decode('utf-8', errors='ignore')
        except Exception:
            return ''

    def extract_text_from_path(p: Path) -> str:
        name = p.name.lower()
        try:
            if name.endswith(('.txt', '.md', '.markdown')):
                return p.read_text(encoding='utf-8', errors='ignore')
            if name.endswith('.json'):
                import json as _json
                data = _json.loads(p.read_text(encoding='utf-8', errors='ignore'))
                if isinstance(data, list):
                    return "\n\n".join([str(x) for x in data if str(x).strip()])
                if isinstance(data, dict) and 'documents' in data:
                    return "\n\n".join([str(x) for x in data['documents'] if str(x).strip()])
                return ""
            if name.endswith('.pdf'):
                from pypdf import PdfReader
                reader = PdfReader(str(p))
                pages = [pg.extract_text() or '' for pg in reader.pages]
                return "\n\n".join(pages)
            if name.endswith('.docx'):
                from docx import Document
                doc = Document(str(p))
                paras = [q.text for q in doc.paragraphs if q.text and q.text.strip()]
                return "\n".join(paras)
            # Fallback
            return p.read_text(encoding='utf-8', errors='ignore')
        except Exception:
            app.logger.exception(f"Failed to parse {p}")
            return ''

    # ---- Jinja filters ----
    def _highlight(text: str, query: str) -> str:
        try:
            import re
            if not text or not query:
                return text or ""
            # Escape regex special chars in query, split into words, highlight each
            words = [w for w in re.split(r"\s+", query.strip()) if w]
            if not words:
                return text
            def repl(m):
                return f"<mark class=hl>{m.group(0)}</mark>"
            out = text
            for w in words:
                pat = re.compile(re.escape(w), flags=re.IGNORECASE)
                out = pat.sub(repl, out)
            return out
        except Exception:
            return text or ""
    app.jinja_env.filters['hl'] = _highlight

    @app.route("/", methods=["GET"]) 
    def home():
        # Default to chat UI; the search UI is still available via POST /search from index.html
        return redirect(url_for("chat_home"))

    @app.route("/search", methods=["POST"]) 
    def search():
        query = request.form.get("q", "").strip()
        try:
            top_k = int(request.form.get("top_k", "10"))
        except Exception:
            top_k = 10

        if not query:
            flash("Enter a query to search.", "warning")
            return redirect(url_for("home"))

        try:
            result = system.search(query, top_k=top_k)
            stats = system.get_system_info()
            return render_template("index.html", query=query, top_k=top_k, result=result, stats=stats)
        except Exception as e:
            app.logger.exception("Search failed")
            flash(str(e), "danger")
            return redirect(url_for("home"))

    @app.route("/api/search", methods=["GET"]) 
    def api_search():
        query = request.args.get("q", "").strip()
        top_k = request.args.get("top_k", type=int, default=10)
        if not query:
            return jsonify({"error": "Missing q"}), 400
        try:
            result = system.search(query, top_k=top_k)
            return jsonify(result)
        except Exception as e:
            app.logger.exception("API search failed")
            return jsonify({"error": str(e)}), 500

    @app.route("/add", methods=["POST"]) 
    def add_documents():
        # Add documents via textarea or file upload (supports txt/json/pdf/docx/md)
        docs_text = request.form.get("docs", "").strip()
        uploads = request.files.getlist("file")
        docs = []
        if docs_text:
            docs.extend([d.strip() for d in docs_text.split("\n") if d.strip()])
        upload_summaries = []
        for uploaded in uploads:
            if uploaded and uploaded.filename:
                text = extract_text_from_upload(uploaded)
                if text:
                    chunks = chunk_text(text)
                    docs.extend(chunks)
                    upload_summaries.append({
                        "filename": uploaded.filename,
                        "hash": text_hash(text),
                        "chunks": len(chunks),
                        "bytes": len(text.encode('utf-8', errors='ignore')),
                        "source": "upload",
                        "time": int(time.time()),
                    })
        # If request comes from browser form, prefer redirect with flash
        wants_html = request.accept_mimetypes.accept_html and not request.accept_mimetypes.accept_json
        if not docs:
            if wants_html:
                flash("No documents provided", "warning")
                return redirect(request.referrer or url_for("home"))
            return jsonify({"added": 0, "message": "No documents provided"}), 400

        try:
            count = system.add_documents(docs, source="webui")
            if upload_summaries:
                m = load_manifest()
                m_files = m.get('files', [])
                # Avoid duplicates by hash
                existing_hashes = {f.get('hash') for f in m_files}
                for us in upload_summaries:
                    if us['hash'] and us['hash'] not in existing_hashes:
                        m_files.append(us)
                        existing_hashes.add(us['hash'])
                m['files'] = m_files
                save_manifest(m)
            if wants_html:
                flash(f"Embedded {count} document chunks", "success")
                return redirect(request.referrer or url_for("home"))
            return jsonify({"added": count, "total": len(system.doc_manager.get_documents())})
        except Exception as e:
            app.logger.exception("Failed to add documents")
            if wants_html:
                flash(str(e), "danger")
                return redirect(request.referrer or url_for("home"))
            return jsonify({"error": str(e)}), 500

    # ---- Chat-like Q&A ----
    @app.route("/chat", methods=["GET"]) 
    def chat_home():
        history = session.get('chat_history', [])
        stats = system.get_system_info()
        return render_template("chat.html", history=history, stats=stats)

    @app.route("/chat/send", methods=["POST"]) 
    def chat_send():
        q = request.form.get('message', '').strip()
        top_k = request.form.get('top_k', type=int, default=5)
        if not q:
            return redirect(url_for('chat_home'))
        try:
            result = system.search(q, top_k=top_k)
            # Compose answer from the final stage top-1 passage (no LLM)
            final = (result or {}).get('results') or []
            if final:
                best = final[0]
                passage = (best.get('document') or '').strip()
                # Prefer stage3_score; fallback to stage2/1
                score = best.get('stage3_score')
                if score is None:
                    score = best.get('stage2_score')
                if score is None:
                    score = best.get('score')
                footer = f"\n\n[stage3_score: {score}]" if score is not None else ''
                answer = passage[:4000] + footer
            else:
                answer = "No results."
            history = session.get('chat_history', [])
            history.append({"role": "user", "content": q})
            history.append({"role": "assistant", "content": answer})
            session['chat_history'] = history[-50:]  # keep last 50 messages
        except Exception as e:
            app.logger.exception('Chat search failed')
            flash(str(e), 'danger')
        return redirect(url_for('chat_home'))

    @app.route("/api/clear", methods=["POST"]) 
    def api_clear():
        try:
            system.clear_all_data()
            session['chat_history'] = []
            return jsonify({"status": "cleared"})
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    @app.route("/api/stats", methods=["GET"]) 
    def api_stats():
        try:
            info = system.get_system_info()
            return jsonify(info)
        except Exception as e:
            return jsonify({"error": str(e)}), 500

    # --- Embed management UI and APIs ---
    @app.route("/embed", methods=["GET"]) 
    def embed_page():
        statuses = repo_documents_status()
        manifest = load_manifest()
        stats = system.get_system_info()
        return render_template("embed.html", statuses=statuses, manifest=manifest, stats=stats)

    @app.route("/embed/run", methods=["POST"]) 
    def embed_run():
        action = request.form.get("action", "").strip()
        chunk_size = request.form.get("chunk_size", type=int, default=1000)
        overlap = request.form.get("overlap", type=int, default=200)
        try:
            if action == "upload":
                uploads = request.files.getlist("file")
                docs = []
                new_entries = []
                for uploaded in uploads:
                    if uploaded and uploaded.filename:
                        text = extract_text_from_upload(uploaded)
                        if text:
                            chunks = chunk_text(text, chunk_size, overlap)
                            docs.extend(chunks)
                            new_entries.append({
                                "filename": uploaded.filename,
                                "hash": text_hash(text),
                                "chunks": len(chunks),
                                "bytes": len(text.encode('utf-8', errors='ignore')),
                                "source": "upload",
                                "time": int(time.time()),
                            })
                added = system.add_documents(docs, source="embed-upload") if docs else 0
                if new_entries:
                    m = load_manifest()
                    m_files = m.get('files', [])
                    existing = {f.get('hash') for f in m_files}
                    for e in new_entries:
                        if e['hash'] and e['hash'] not in existing:
                            m_files.append(e)
                            existing.add(e['hash'])
                    m['files'] = m_files
                    save_manifest(m)
                flash(f"Embedded {added} chunks from uploads", "success")
                return redirect(url_for('embed_page'))
            elif action == "repo_all":
                statuses = repo_documents_status()
                to_embed = [s for s in statuses if not s.get('embedded') and s.get('hash')]
                total_chunks = 0
                new_entries = []
                for s in to_embed:
                    p = PROJECT_ROOT / s['path']
                    text = extract_text_from_path(p)
                    if text:
                        chunks = chunk_text(text, chunk_size, overlap)
                        total_chunks += len(chunks)
                        system.add_documents(chunks, source="repo-documents")
                        new_entries.append({
                            "path": s['path'],
                            "hash": s['hash'],
                            "chunks": len(chunks),
                            "bytes": len(text.encode('utf-8', errors='ignore')),
                            "source": "repo-documents",
                            "time": int(time.time()),
                        })
                if new_entries:
                    m = load_manifest()
                    m_files = m.get('files', [])
                    existing = {f.get('hash') for f in m_files}
                    for e in new_entries:
                        if e['hash'] and e['hash'] not in existing:
                            m_files.append(e)
                            existing.add(e['hash'])
                    m['files'] = m_files
                    save_manifest(m)
                flash(f"Embedded {total_chunks} chunks from documents/", "success")
                return redirect(url_for('embed_page'))
            elif action == "repo_file":
                rel_path = request.form.get('path', '')
                if not rel_path:
                    flash("Missing path", "warning")
                    return redirect(url_for('embed_page'))
                p = PROJECT_ROOT / rel_path
                text = extract_text_from_path(p)
                if not text:
                    flash("No text parsed from file", "warning")
                    return redirect(url_for('embed_page'))
                chunks = chunk_text(text, chunk_size, overlap)
                system.add_documents(chunks, source="repo-documents")
                m = load_manifest()
                m_files = m.get('files', [])
                h = text_hash(text)
                if h and h not in {f.get('hash') for f in m_files}:
                    m_files.append({
                        "path": str(Path(rel_path)),
                        "hash": h,
                        "chunks": len(chunks),
                        "bytes": len(text.encode('utf-8', errors='ignore')),
                        "source": "repo-documents",
                        "time": int(time.time()),
                    })
                    m['files'] = m_files
                    save_manifest(m)
                flash(f"Embedded {len(chunks)} chunks from {rel_path}", "success")
                return redirect(url_for('embed_page'))
            else:
                flash("Unknown action", "danger")
                return redirect(url_for('embed_page'))
        except Exception as e:
            app.logger.exception('Embed run failed')
            flash(str(e), 'danger')
            return redirect(url_for('embed_page'))

    @app.route("/api/embedded", methods=["GET"]) 
    def api_embedded():
        return jsonify(load_manifest())

    @app.route("/api/documents-status", methods=["GET"]) 
    def api_documents_status():
        return jsonify({"statuses": repo_documents_status()})

    # Auto-ingest disabled by request; documents are only added via /add or CLI

    return app


def main():
    """Entry point for the ``tristage-webui`` GUI script.

    Sets the CWD to the repo root so relative paths (./models, ./faiss_index)
    resolve correctly, then launches the Flask dev server.
    """
    import os
    os.chdir(PROJECT_ROOT)
    app = create_app()
    app.run(host="127.0.0.1", port=5051, debug=False)


if __name__ == "__main__":
    main()
